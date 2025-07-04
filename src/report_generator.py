from docx import Document
from datetime import datetime

from main_window import MyMainWindow

def fill_template(template_path, output_path, context):
    doc = Document(template_path)
    
    # Замена текста в параграфах
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    
    # Замена текста в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
    
    doc.save(output_path)
    print(f"Отчет сохранен: {output_path}")

obj = MyMainWindow()

# Данные для отчета
context = {
    "{REPORT_DATE}": datetime.now().strftime("%d.%m.%Y"),
    "{START_DATE}": obj.fileModel.fileName(obj.folderView.currentIndex()),
    "{END_DATE}": obj.chart,
    "{PPN_1}": "1 245 000",
    "{PPN_2}": "42",
    "{PPN_3}": "29 643",
    "{CONCLUSION}": "Продажи выросли на 15% по сравнению с предыдущим периодом.",
    "{PBPI_1}": "Иванов И.И.",
    "{PBPI_2}": "Менеджер по продажам",
    "{PBPI_3}": "Ноутбук X",
    "{TBPI_1}": "10",
    "{TBPI_2}": "500 000",
    "{TBPI_3}": "Смартфон Y",
    
}

fill_template("templates/template.docx", "reports/report_filled.docx", context)